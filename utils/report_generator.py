# utils/report_generator.py — генерация отчета в Word

import streamlit as st
from docx import Document
from docx.shared import Inches
import os

"""
Генерация отчёта .docx с визуализациями и комментариями.
"""


def export_to_word(output_file="edu_monitor_report.docx", standard_text=None) -> bool:
    """Создаёт и сохраняет отчёт в формате Word"""
    try:
        saved_plots = st.session_state.get("saved_plots", [])
        plot_descriptions = st.session_state.get("plot_descriptions", [])

        if not saved_plots:
            st.warning("⚠️ Нет графиков для экспорта.")
            return False

        doc = Document()
        doc.add_heading("Отчёт по анализу данных мониторинга вузов", 0)

        doc.add_paragraph(
            standard_text or (
                "Отчёт включает результаты анализа, включая визуализации, корреляции, кластеризацию и предсказания."
            )
        )

        doc.add_heading("Визуализации", level=1)
        for path, description in zip(saved_plots, plot_descriptions):
            if os.path.exists(path):
                doc.add_paragraph(description)
                doc.add_picture(path, width=Inches(6))

        doc.save(output_file)
        st.success(f"📄 Отчёт сохранён: {output_file}")
        return True

    except Exception as e:
        st.error(f"❌ Ошибка создания отчёта: {e}")
        return False
