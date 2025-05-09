import streamlit as st
import pandas as pd
import joblib
import os
from docx import Document
from docx.shared import Inches


saved_plots = []
plot_descriptions = []
@st.cache_data
def save_dataframe(df: pd.DataFrame, filename: str) -> bool:
    """Сохраняет DataFrame в CSV-файл."""
    try:
        df.to_csv(filename, index=False)
        st.success(f"💾 Данные сохранены: {filename}")
        return True
    except Exception as e:
        st.error(f"❌ Ошибка при сохранении файла: {e}")
        return False


def load_model(path: str):
    """Загружает модель из файла с помощью joblib."""
    try:
        model = joblib.load(path)
        st.success(f"📥 Модель загружена: {path}")
        return model
    except Exception as e:
        st.error(f"❌ Ошибка при загрузке модели: {e}")
        return None


def save_model(model, path: str) -> None:
    """Сохраняет модель в файл с помощью joblib."""
    try:
        joblib.dump(model, path)
        st.success(f"📤 Модель сохранена: {path}")
    except Exception as e:
        st.error(f"❌ Ошибка при сохранении модели: {e}")


def export_to_word(output_file="edu_monitor_report.docx", standard_text=None):
    try:
        doc = Document()
        doc.add_heading("Отчёт по анализу данных мониторинга вузов", 0)

        doc.add_paragraph(
            standard_text or (
                "Отчёт включает результаты анализа, включая визуализации, корреляции, кластеризацию и предсказания."
            )
        )

        doc.add_heading("Визуализации", level=1)
        for path, description in zip(saved_plots, plot_descriptions):
            if os.path.exists(path):
                doc.add_paragraph(description)
                doc.add_picture(path, width=Inches(6))
        doc.save(output_file)
        st.success(f"Отчёт сохранён как {output_file}")
        return True
    except Exception as e:
        st.error(f"Ошибка создания отчёта: {e}")
        return False


def clear_saved_plots():
    try:
        for file in os.listdir("plots"):
            os.remove(os.path.join("plots", file))
        saved_plots.clear()
        plot_descriptions.clear()
        st.info("Очистка графиков выполнена")
        return True
    except Exception as e:
        st.error(f"Ошибка при очистке графиков: {e}")
        return False


def detect_outliers(df: pd.DataFrame, column: str, iqr_factor: float = 1.5):
    try:
        q1 = df[column].quantile(0.25)
        q3 = df[column].quantile(0.75)
        iqr = q3 - q1
        return (df[column] < q1 - iqr_factor * iqr) | (df[column] > q3 + iqr_factor * iqr)
    except Exception as e:
        st.error(f"Ошибка при поиске выбросов: {e}")
        return None
